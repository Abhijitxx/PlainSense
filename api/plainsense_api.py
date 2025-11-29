"""
PlainSense API Module
=====================

Unified API for the PlainSense document simplification system.
Provides endpoints for both Legal and Medical document processing.

Features:
1. Legal Document Processing
   - OCR text extraction
   - Clause segmentation
   - Multi-language simplification (English, Hindi, Tamil)
   - Risk detection and warnings
   - Entity preservation

2. Medical Document Processing
   - OCR with table extraction
   - Lab report parsing
   - Medical term explanations
   - Risk assessment
   - Multi-language output

Author: PlainSense Team
Date: November 2025
"""

import os
import sys
import time
from typing import Dict, List, Optional, Union
from dataclasses import dataclass, field, asdict
from enum import Enum

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import project modules
from core.llm_simplifier import LLMSimplifier, ClauseResult, DocumentResult, RiskLevel
from core.clause_segmenter import ClauseSegmenter
from medical.medical_dictionary import MedicalDictionary

# Try to import OCR modules (may not be available on all systems)
try:
    from ocr.ocr_pipeline import OCRPipeline
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR module not available")

try:
    from medical.medical_report_parser import MedicalReportParser
    MEDICAL_PARSER_AVAILABLE = True
except ImportError:
    MEDICAL_PARSER_AVAILABLE = False
    print("⚠️ Medical report parser not available")

# Import OCR text corrector
try:
    from core.text_corrector import OCRTextCorrector, LawReference
    TEXT_CORRECTOR_AVAILABLE = True
except ImportError:
    TEXT_CORRECTOR_AVAILABLE = False
    print("⚠️ Text corrector not available")


class DocumentType(str, Enum):
    LEGAL = "legal"
    MEDICAL = "medical"


@dataclass
class ProcessingResult:
    """Unified result structure for document processing"""
    success: bool
    document_type: str
    clauses: List[Dict] = field(default_factory=list)
    summary: Dict = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)
    processing_time: float = 0.0
    
    # Medical-specific fields
    lab_results: List[Dict] = field(default_factory=list)
    medical_explanations: List[Dict] = field(default_factory=list)
    
    # Law explanations for popup display
    law_references: List[Dict] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)


class PlainSenseAPI:
    """
    Unified API for PlainSense document simplification.
    
    Usage:
        api = PlainSenseAPI()
        
        # For legal documents
        result = api.process_legal_document(text_or_file_path)
        
        # For medical documents  
        result = api.process_medical_document(text_or_file_path)
    """
    
    def __init__(self, domain: str = None, device: str = "auto"):
        """
        Initialize the PlainSense API.
        
        Args:
            domain: 'legal', 'medical', or None for both
            device: Device to use ('cpu', 'cuda', 'auto')
        """
        print("=" * 60)
        print("  INITIALIZING PLAINSENSE API")
        print("=" * 60)
        
        self.domain = domain
        
        # Initialize core components with domain-specific loading
        print("\n📝 Loading LLM Simplifier...")
        self.simplifier = LLMSimplifier(device=device, domain=domain)
        
        print("📋 Loading Clause Segmenter...")
        self.segmenter = ClauseSegmenter()
        
        # Initialize OCR text corrector with aggressive mode
        self.text_corrector = None
        if TEXT_CORRECTOR_AVAILABLE:
            print("🔧 Loading OCR Text Corrector (aggressive mode)...")
            self.text_corrector = OCRTextCorrector(use_ml=True, aggressive_mode=True)
            
            # Skip LLM for law explanations - use basic explanations for speed
            # The LLM call per law reference is too slow
            # if self.simplifier and self.simplifier.models_loaded:
            #     print("   🔗 Connecting LLM for dynamic law explanations...")
            #     self.text_corrector.set_llm_callback(self._explain_law_with_llm)
        
        # Only load medical dictionary if needed
        self.med_dictionary = None
        if domain is None or domain == 'medical':
            print("📚 Loading Medical Dictionary...")
            self.med_dictionary = MedicalDictionary()
        
        # Initialize optional components
        self.ocr = None
        if OCR_AVAILABLE:
            print("🔍 Loading OCR Pipeline...")
            self.ocr = OCRPipeline()
        
        self.medical_parser = None
        if MEDICAL_PARSER_AVAILABLE and (domain is None or domain == 'medical'):
            print("🏥 Loading Medical Report Parser...")
            self.medical_parser = MedicalReportParser()
        
        print("\n✅ PlainSense API initialized successfully!")
        print("=" * 60)
    
    def _explain_law_with_llm(self, prompt: str) -> str:
        """
        Use the LLM to generate explanations for law references.
        This is called by the DynamicLawExplainer.
        
        Args:
            prompt: The prompt describing the law to explain
            
        Returns:
            LLM-generated explanation
        """
        if not self.simplifier or not self.simplifier.models_loaded:
            return ""
        
        try:
            # Use the simplifier's model to generate explanation
            # We'll use a simple approach - encode prompt and decode response
            inputs = self.simplifier.simplifier_tokenizer(
                prompt,
                max_length=512,
                truncation=True,
                return_tensors="pt"
            )
            
            if hasattr(self.simplifier, 'device'):
                inputs = {k: v.to(self.simplifier.device) for k, v in inputs.items()}
            
            outputs = self.simplifier.simplifier_model.generate(
                inputs['input_ids'],
                attention_mask=inputs.get('attention_mask'),
                max_length=200,
                num_beams=2,
                early_stopping=True,
                no_repeat_ngram_size=2
            )
            
            explanation = self.simplifier.simplifier_tokenizer.decode(
                outputs[0], 
                skip_special_tokens=True
            ).strip()
            
            return explanation if len(explanation) > 30 else ""
            
        except Exception as e:
            print(f"   ⚠️ LLM law explanation failed: {e}")
            return ""
    
    def process_legal_document(
        self,
        input_source: Union[str, List[str]],
        include_translations: bool = False,
        translation_lang: str = None,
        output_style: str = "all"
    ) -> ProcessingResult:
        """
        Process a legal document (rental agreement, contract, etc.)
        
        Args:
            input_source: File path, raw text, or list of clause texts
            include_translations: Whether to generate translations (loads models on demand)
            translation_lang: Specific language ('hi', 'ta') or None for all
            output_style: Which output styles to generate
            
        Returns:
            ProcessingResult with simplified clauses and risk analysis
        """
        start_time = time.time()
        errors = []
        
        # Step 1: Get text from input
        text = ""
        clauses = []
        
        if isinstance(input_source, list):
            # Already segmented clauses
            clauses = input_source
        elif os.path.isfile(input_source):
            # File path - extract text based on file type
            file_ext = os.path.splitext(input_source)[1].lower()
            
            try:
                if file_ext == '.docx':
                    # Handle DOCX files
                    try:
                        from docx import Document
                        doc = Document(input_source)
                        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    except ImportError:
                        errors.append("python-docx not installed for DOCX support")
                elif file_ext == '.pdf':
                    # Handle PDF files
                    if self.ocr:
                        text = self.ocr.extract_text(input_source)
                    else:
                        try:
                            from pypdf import PdfReader
                            reader = PdfReader(input_source)
                            text = '\n'.join([page.extract_text() or '' for page in reader.pages])
                        except ImportError:
                            errors.append("pypdf not installed for PDF support")
                elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    # Handle image files with OCR
                    if self.ocr:
                        text = self.ocr.extract_text(input_source)
                    else:
                        errors.append("OCR not available for image files")
                else:
                    # Try reading as text file
                    with open(input_source, 'r', encoding='utf-8') as f:
                        text = f.read()
            except Exception as e:
                errors.append(f"File read failed: {str(e)}")
        else:
            # Raw text input
            text = input_source
        
        # Step 1.5: Apply OCR text correction to fix common OCR errors
        # Also detects law references for popup explanations
        law_references = []
        if text and self.text_corrector:
            print(f"[DEBUG API] Applying OCR text correction (aggressive mode)...")
            original_len = len(text)
            text = self.text_corrector.correct_legal_document(text)
            print(f"[DEBUG API] Text corrected: {original_len} -> {len(text)} chars")
            
            # Get detected law references for popup explanations
            law_references = self.text_corrector.get_law_explanations()
            if law_references:
                print(f"[DEBUG API] Detected {len(law_references)} law references for explanations")
        
        # Step 2: Segment into clauses
        if not clauses and text:
            try:
                print(f"[DEBUG API] Segmenting text of length {len(text)}")
                segment_result = self.segmenter.segment(text, doc_type='legal')
                print(f"[DEBUG API] Segmentation result: {segment_result.get('total_clauses', 0)} clauses")
                print(f"[DEBUG API] Segmentation method: {segment_result.get('segmentation_method', 'unknown')}")
                # Extract clauses from result dict - get the content of each clause
                if isinstance(segment_result, dict) and 'clauses' in segment_result:
                    clauses = [c['content'] for c in segment_result['clauses']]
                    print(f"[DEBUG API] Extracted {len(clauses)} clause contents")
                else:
                    clauses = [text]
                    print(f"[DEBUG API] Fallback: using entire text as one clause")
            except Exception as e:
                errors.append(f"Segmentation failed: {str(e)}")
                clauses = [text]  # Treat entire text as one clause
                print(f"[DEBUG API] Segmentation error: {e}")
        
        # Step 3: Simplify each clause
        simplified_clauses = []
        risk_counts = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'NONE': 0}
        
        for i, clause in enumerate(clauses):
            print(f"   Processing clause {i+1}/{len(clauses)}...", end="\r")
            
            try:
                result = self.simplifier.simplify_clause(
                    clause,
                    domain='legal',
                    include_translations=include_translations
                )
                
                clause_dict = result.to_dict()
                clause_dict['clause_number'] = i + 1
                simplified_clauses.append(clause_dict)
                
                # Count risk levels
                risk_counts[result.risk_level.name] += 1
                
            except Exception as e:
                errors.append(f"Clause {i+1} processing failed: {str(e)}")
                simplified_clauses.append({
                    'clause_number': i + 1,
                    'original': clause if isinstance(clause, str) else str(clause),
                    'error': str(e)
                })
        
        print(" " * 50)  # Clear progress line
        
        # Step 4: Generate summary
        summary = {
            'total_clauses': len(clauses),
            'risk_breakdown': risk_counts,
            'high_risk_count': risk_counts['CRITICAL'] + risk_counts['HIGH'],
            'overall_risk': self._calculate_overall_risk(risk_counts),
            'recommendations': self._generate_legal_recommendations(simplified_clauses)
        }
        
        processing_time = time.time() - start_time
        
        return ProcessingResult(
            success=len(errors) == 0,
            document_type='legal',
            clauses=simplified_clauses,
            summary=summary,
            errors=errors,
            processing_time=processing_time,
            law_references=law_references  # Include law references for popup explanations
        )
    
    def process_medical_document(
        self,
        input_source: Union[str, List[str]],
        include_translations: bool = None,
        include_dictionary: bool = True
    ) -> ProcessingResult:
        """
        Process a medical document (lab report, prescription, etc.)
        
        Args:
            input_source: File path, raw text, or list of result texts
            include_translations: Override for translation setting
            include_dictionary: Whether to include medical dictionary explanations
            
        Returns:
            ProcessingResult with simplified results and medical explanations
        """
        start_time = time.time()
        errors = []
        
        if include_translations is None:
            include_translations = self.enable_translations
        
        # Step 1: Get text from input
        text = ""
        sections = []
        lab_results = []
        
        if isinstance(input_source, list):
            sections = input_source
        elif os.path.isfile(input_source):
            # File path - extract text based on file type
            file_ext = os.path.splitext(input_source)[1].lower()
            
            try:
                if file_ext == '.docx':
                    # Handle DOCX files
                    try:
                        from docx import Document
                        doc = Document(input_source)
                        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    except ImportError:
                        errors.append("python-docx not installed for DOCX support")
                elif file_ext == '.pdf':
                    # Handle PDF files
                    if self.ocr:
                        text = self.ocr.extract_text(input_source, detect_tables=True)
                    else:
                        try:
                            from pypdf import PdfReader
                            reader = PdfReader(input_source)
                            text = '\n'.join([page.extract_text() or '' for page in reader.pages])
                        except ImportError:
                            errors.append("pypdf not installed for PDF support")
                elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    # Handle image files with OCR
                    if self.ocr:
                        text = self.ocr.extract_text(input_source, detect_tables=True)
                    else:
                        errors.append("OCR not available for image files")
                else:
                    # Try reading as text file
                    with open(input_source, 'r', encoding='utf-8') as f:
                        text = f.read()
            except Exception as e:
                errors.append(f"File read failed: {str(e)}")
            
            # Try to parse structured lab results
            if self.medical_parser and text:
                try:
                    lab_results = self.medical_parser.parse(text)
                except Exception as e:
                    errors.append(f"Lab parsing failed: {str(e)}")
        else:
            text = input_source
        
        # Step 2: Segment into sections
        if not sections and text:
            try:
                segment_result = self.segmenter.segment(text, doc_type='medical')
                # Extract sections from result dict - get the content of each section
                if isinstance(segment_result, dict) and 'clauses' in segment_result:
                    sections = [c['content'] for c in segment_result['clauses']]
                else:
                    sections = [text]
            except Exception as e:
                errors.append(f"Segmentation failed: {str(e)}")
                sections = [text]
        
        # Step 3: Simplify each section
        simplified_sections = []
        risk_counts = {'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'NONE': 0}
        
        for i, section in enumerate(sections):
            print(f"   Processing section {i+1}/{len(sections)}...", end="\r")
            
            try:
                result = self.simplifier.simplify_clause(
                    section,
                    domain='medical',
                    include_translations=include_translations
                )
                
                section_dict = result.to_dict()
                section_dict['section_number'] = i + 1
                simplified_sections.append(section_dict)
                
                risk_counts[result.risk_level.name] += 1
                
            except Exception as e:
                errors.append(f"Section {i+1} processing failed: {str(e)}")
                simplified_sections.append({
                    'section_number': i + 1,
                    'original': section if isinstance(section, str) else str(section),
                    'error': str(e)
                })
        
        print(" " * 50)
        
        # Step 4: Get medical dictionary explanations
        medical_explanations = []
        if include_dictionary and lab_results:
            medical_explanations = self.med_dictionary.interpret_report(lab_results)
        
        # Step 5: Generate summary
        summary = {
            'total_sections': len(sections),
            'lab_tests_found': len(lab_results),
            'risk_breakdown': risk_counts,
            'abnormal_results': risk_counts['CRITICAL'] + risk_counts['HIGH'],
            'overall_health_risk': self._calculate_overall_risk(risk_counts),
            'recommendations': self._generate_medical_recommendations(simplified_sections, medical_explanations)
        }
        
        processing_time = time.time() - start_time
        
        return ProcessingResult(
            success=len(errors) == 0,
            document_type='medical',
            clauses=simplified_sections,
            summary=summary,
            errors=errors,
            processing_time=processing_time,
            lab_results=[asdict(r) if hasattr(r, '__dataclass_fields__') else r for r in lab_results],
            medical_explanations=medical_explanations
        )
    
    def explain_medical_term(self, term: str) -> Dict:
        """
        Get a simple explanation for a medical term.
        
        Args:
            term: Medical term (test name or condition)
            
        Returns:
            Dictionary with explanation
        """
        # Try as test first
        test_info = self.med_dictionary.lookup_test(term)
        if test_info:
            return {
                'found': True,
                'type': 'test',
                'term': test_info.term,
                'simple_name': test_info.simple_name,
                'description': test_info.description,
                'normal_range': test_info.normal_range,
                'high_meaning': test_info.high_meaning,
                'low_meaning': test_info.low_meaning
            }
        
        # Try as condition
        condition_info = self.med_dictionary.lookup_condition(term)
        if condition_info:
            return {
                'found': True,
                'type': 'condition',
                'term': condition_info.term,
                'simple_name': condition_info.simple_name,
                'description': condition_info.description,
                'symptoms': condition_info.symptoms
            }
        
        return {
            'found': False,
            'term': term,
            'message': f"'{term}' not found in medical dictionary"
        }
    
    def interpret_lab_result(
        self,
        test_name: str,
        value: float,
        unit: str = ""
    ) -> Dict:
        """
        Interpret a single lab test result.
        
        Args:
            test_name: Name of the test
            value: Test value
            unit: Unit of measurement
            
        Returns:
            Dictionary with interpretation
        """
        return self.med_dictionary.explain_test_result(test_name, value, unit)
    
    def _calculate_overall_risk(self, risk_counts: Dict[str, int]) -> str:
        """Calculate overall risk level from counts"""
        if risk_counts.get('CRITICAL', 0) > 0:
            return 'CRITICAL'
        if risk_counts.get('HIGH', 0) >= 2:
            return 'HIGH'
        if risk_counts.get('HIGH', 0) >= 1 or risk_counts.get('MEDIUM', 0) >= 3:
            return 'MEDIUM'
        if risk_counts.get('MEDIUM', 0) >= 1:
            return 'LOW'
        return 'NONE'
    
    def _generate_legal_recommendations(self, clauses: List[Dict]) -> List[str]:
        """Generate recommendations for legal documents"""
        recommendations = []
        
        high_risk_clauses = [c for c in clauses if c.get('risk', {}).get('level') in ['CRITICAL', 'HIGH']]
        
        if high_risk_clauses:
            recommendations.append(
                f"⚠️ Found {len(high_risk_clauses)} high-risk clauses that may be unfavorable. Consider negotiating these terms."
            )
        
        # Check for common issues
        for clause in clauses:
            original = clause.get('original', '').lower()
            if 'forfeit' in original and 'deposit' in original:
                recommendations.append(
                    "💰 Document contains forfeiture clauses for deposits. Ensure conditions are fair and clearly defined."
                )
                break
        
        for clause in clauses:
            original = clause.get('original', '').lower()
            if 'without notice' in original:
                recommendations.append(
                    "📋 Some clauses allow actions without notice. Request reasonable notice periods."
                )
                break
        
        if not recommendations:
            recommendations.append("✅ Document appears to have reasonable terms. Review all clauses carefully before signing.")
        
        return recommendations
    
    def _generate_medical_recommendations(
        self,
        sections: List[Dict],
        explanations: List[Dict]
    ) -> List[str]:
        """Generate recommendations for medical documents"""
        recommendations = []
        
        # Check for abnormal results
        abnormal = [e for e in explanations if e.get('status') in ['high', 'low']]
        
        if abnormal:
            recommendations.append(
                f"🔬 Found {len(abnormal)} abnormal test results. Please discuss these with your doctor."
            )
            
            # Specific recommendations based on tests
            for result in abnormal:
                test = result.get('test', '').lower()
                status = result.get('status', '')
                
                if 'hemoglobin' in test and status == 'low':
                    recommendations.append(
                        "🩸 Low hemoglobin detected. You may have anemia. Consider iron-rich foods and consult a doctor."
                    )
                elif 'glucose' in test and status == 'high':
                    recommendations.append(
                        "🍬 High blood sugar detected. Monitor your diet and follow up with diabetes screening."
                    )
                elif 'creatinine' in test and status == 'high':
                    recommendations.append(
                        "🏥 Elevated creatinine suggests kidney function concerns. Stay hydrated and see a nephrologist."
                    )
        else:
            recommendations.append("✅ Test results appear within normal ranges. Continue with regular health checkups.")
        
        return recommendations[:5]  # Limit to 5 recommendations


# Convenience functions for direct use
def simplify_legal(text: str, include_translations: bool = True) -> ProcessingResult:
    """Quick function to simplify legal text"""
    api = PlainSenseAPI(enable_translations=include_translations)
    return api.process_legal_document(text)


def simplify_medical(text: str, include_translations: bool = True) -> ProcessingResult:
    """Quick function to simplify medical text"""
    api = PlainSenseAPI(enable_translations=include_translations)
    return api.process_medical_document(text)


if __name__ == "__main__":
    # Quick test
    print("\n" + "=" * 60)
    print("  PLAINSENSE API TEST")
    print("=" * 60)
    
    api = PlainSenseAPI(enable_translations=False)
    
    # Test legal
    print("\n📜 Testing Legal Processing...")
    legal_text = """
    The Tenant shall pay a security deposit of Rs. 50,000 which will be 
    forfeited if they vacate before 6 months. Either party may terminate 
    with 30 days written notice.
    """
    
    result = api.process_legal_document(legal_text)
    print(f"   Clauses processed: {result.summary['total_clauses']}")
    print(f"   Overall risk: {result.summary['overall_risk']}")
    print(f"   Time: {result.processing_time:.2f}s")
    
    # Test medical
    print("\n🏥 Testing Medical Processing...")
    medical_text = """
    Hemoglobin: 10.5 g/dL (Normal: 12-17)
    Fasting Blood Sugar: 180 mg/dL (Normal: 70-100)
    Creatinine: 1.1 mg/dL (Normal: 0.7-1.3)
    """
    
    result = api.process_medical_document(medical_text)
    print(f"   Sections processed: {result.summary['total_sections']}")
    print(f"   Overall risk: {result.summary['overall_health_risk']}")
    print(f"   Time: {result.processing_time:.2f}s")
    
    # Test medical dictionary
    print("\n📚 Testing Medical Dictionary...")
    explanation = api.explain_medical_term("hemoglobin")
    print(f"   Hemoglobin: {explanation.get('simple_name', 'N/A')}")
    
    print("\n" + "=" * 60)
    print("  API TEST COMPLETE")
    print("=" * 60)
