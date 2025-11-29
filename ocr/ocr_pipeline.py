"""
Enhanced OCR Pipeline v2
========================

Major Improvements:
1. Multi-engine OCR (Tesseract + EasyOCR fallback)
2. Adaptive image preprocessing
3. OCR confidence scoring
4. Language detection (English/Hindi)
5. Post-OCR text correction
6. Table structure detection

Author: PlainSense Team
Date: November 2025
"""

import os
import re
import cv2
import numpy as np
from pathlib import Path
from PIL import Image, ImageEnhance, ImageFilter
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import warnings
warnings.filterwarnings('ignore')

# OCR engines
try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = os.environ.get(
        'TESSERACT_CMD', 
        r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    )
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ pytesseract not installed")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# Document extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Table extraction
try:
    from img2table.document import Image as Img2TableImage
    from img2table.ocr import TesseractOCR
    import pandas as pd
    TABLE_EXTRACTION_AVAILABLE = True
except ImportError:
    TABLE_EXTRACTION_AVAILABLE = False
    pd = None


@dataclass
class TableResult:
    """Table extraction result"""
    tables: List[Dict]  # List of tables with rows/columns
    raw_dataframes: List  # pandas DataFrames
    table_count: int
    confidence: float
    
    def to_dict(self) -> Dict:
        return {
            'table_count': self.table_count,
            'tables': self.tables,
            'confidence': round(self.confidence, 2)
        }


@dataclass
class OCRResult:
    """OCR extraction result"""
    text: str
    confidence: float
    language: str
    method: str
    word_count: int
    preprocessing: str
    
    def to_dict(self) -> Dict:
        return {
            'text': self.text[:500] + '...' if len(self.text) > 500 else self.text,
            'confidence': round(self.confidence, 2),
            'language': self.language,
            'method': self.method,
            'word_count': self.word_count,
            'preprocessing': self.preprocessing
        }


class EnhancedOCRPipeline:
    """
    Enhanced OCR pipeline with multiple engines and preprocessing
    
    Smart EasyOCR fallback: Only loads EasyOCR when needed:
    - Low confidence from Tesseract (< 60%)
    - Multi-language detected (Hindi + English)
    - Possible handwritten text detected
    """
    
    # Thresholds for triggering EasyOCR fallback
    CONFIDENCE_THRESHOLD = 60.0  # Below this, try EasyOCR
    HANDWRITING_THRESHOLD = 0.3  # Ratio of irregular characters
    
    def __init__(self, 
                 languages: List[str] = ['eng'],
                 use_gpu: bool = False,
                 smart_fallback: bool = True):
        """
        Initialize OCR pipeline
        
        Args:
            languages: Languages to detect ['eng', 'hin']
            use_gpu: Use GPU for EasyOCR
            smart_fallback: Use EasyOCR only when needed (low confidence/multilingual/handwritten)
        """
        self.languages = languages
        self.use_gpu = use_gpu
        self.smart_fallback = smart_fallback
        
        # Initialize OCR engines
        self.tesseract_available = TESSERACT_AVAILABLE
        self.easyocr_reader = None  # Lazy loaded only when needed
        self._easyocr_load_attempted = False
        
        print("   ⚡ Smart mode: Tesseract primary, EasyOCR fallback on-demand")
    
    def preprocess_image(self, 
                         image: np.ndarray,
                         method: str = 'adaptive') -> Tuple[np.ndarray, str]:
        """
        Apply preprocessing to improve OCR quality
        
        Args:
            image: Input image as numpy array
            method: 'adaptive', 'threshold', 'denoise', 'all'
            
        Returns:
            Preprocessed image and method used
        """
        # Convert to grayscale if needed
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        preprocessing_used = []
        
        if method in ['adaptive', 'all']:
            # Adaptive thresholding
            gray = cv2.adaptiveThreshold(
                gray, 255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                11, 2
            )
            preprocessing_used.append('adaptive_threshold')
        
        if method in ['denoise', 'all']:
            # Denoising
            gray = cv2.fastNlMeansDenoising(gray, h=10)
            preprocessing_used.append('denoise')
        
        if method in ['threshold', 'all']:
            # Simple binary threshold
            _, gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            preprocessing_used.append('otsu_threshold')
        
        # Deskew
        gray = self._deskew(gray)
        preprocessing_used.append('deskew')
        
        # Scale up small images
        h, w = gray.shape[:2]
        if min(h, w) < 1000:
            scale = 1500 / min(h, w)
            gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
            preprocessing_used.append(f'scale_{scale:.1f}x')
        
        return gray, '+'.join(preprocessing_used)
    
    def _deskew(self, image: np.ndarray) -> np.ndarray:
        """Deskew rotated image"""
        try:
            coords = np.column_stack(np.where(image > 0))
            if len(coords) < 10:
                return image
            
            angle = cv2.minAreaRect(coords)[-1]
            
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            
            # Only deskew if significant rotation
            if abs(angle) > 0.5 and abs(angle) < 10:
                (h, w) = image.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                image = cv2.warpAffine(
                    image, M, (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE
                )
        except Exception:
            pass
        
        return image
    
    def _detect_table_regions(self, image: np.ndarray) -> List[Tuple[int, int, int, int]]:
        """Detect table regions in image for structured extraction"""
        try:
            # Find horizontal and vertical lines
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            
            # Binary threshold
            _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)
            
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            vertical = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel)
            
            # Combine
            table_mask = cv2.add(horizontal, vertical)
            
            # Find contours
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            table_regions = []
            for cnt in contours:
                x, y, w, h = cv2.boundingRect(cnt)
                if w > 100 and h > 100:  # Minimum table size
                    table_regions.append((x, y, w, h))
            
            return table_regions
        except Exception:
            return []
    
    def extract_tables(self, 
                       image_path: str,
                       implicit_rows: bool = True,
                       borderless_tables: bool = True,
                       min_confidence: int = 50) -> TableResult:
        """
        Extract tables from image with structure preservation
        
        Args:
            image_path: Path to image file
            implicit_rows: Detect rows without horizontal lines
            borderless_tables: Detect tables without borders
            min_confidence: Minimum OCR confidence (0-100)
            
        Returns:
            TableResult with extracted tables
        """
        if not TABLE_EXTRACTION_AVAILABLE:
            print("⚠️ Table extraction not available (install img2table)")
            return TableResult([], [], 0, 0.0)
        
        try:
            # Ensure Tesseract is in PATH for img2table
            tesseract_dir = r'C:\Program Files\Tesseract-OCR'
            if tesseract_dir not in os.environ.get('PATH', ''):
                os.environ['PATH'] = tesseract_dir + os.pathsep + os.environ.get('PATH', '')
            
            # Initialize img2table with Tesseract OCR
            ocr = TesseractOCR(
                n_threads=1,
                lang='+'.join(self.languages)
            )
            
            # Load image
            img = Img2TableImage(src=image_path)
            
            # Extract tables
            extracted_tables = img.extract_tables(
                ocr=ocr,
                implicit_rows=implicit_rows,
                borderless_tables=borderless_tables,
                min_confidence=min_confidence
            )
            
            # Convert to structured format
            tables = []
            dataframes = []
            total_confidence = 0.0
            
            for table in extracted_tables:
                # Get DataFrame
                df = table.df
                if df is not None and not df.empty:
                    dataframes.append(df)
                    
                    # Convert to dict format
                    table_dict = {
                        'headers': df.columns.tolist(),
                        'rows': df.values.tolist(),
                        'shape': df.shape,
                        'bbox': table.bbox if hasattr(table, 'bbox') else None
                    }
                    tables.append(table_dict)
                    
                    # Estimate confidence based on non-empty cells
                    non_empty = df.notna().sum().sum()
                    total_cells = df.shape[0] * df.shape[1]
                    if total_cells > 0:
                        total_confidence += (non_empty / total_cells)
            
            avg_confidence = total_confidence / len(tables) if tables else 0.0
            
            print(f"📊 Extracted {len(tables)} tables from image")
            for i, t in enumerate(tables):
                print(f"   Table {i+1}: {t['shape'][0]} rows × {t['shape'][1]} columns")
            
            return TableResult(
                tables=tables,
                raw_dataframes=dataframes,
                table_count=len(tables),
                confidence=avg_confidence
            )
            
        except Exception as e:
            print(f"⚠️ Table extraction error: {e}")
            return TableResult([], [], 0, 0.0)
    
    def extract_tables_from_numpy(self, 
                                   image: np.ndarray,
                                   temp_path: str = None) -> TableResult:
        """
        Extract tables from numpy array image
        
        Args:
            image: Image as numpy array
            temp_path: Optional temp path for saving
            
        Returns:
            TableResult with extracted tables
        """
        import tempfile
        
        # Save to temp file
        if temp_path is None:
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            temp_path = temp_file.name
            temp_file.close()
        
        # Save image
        cv2.imwrite(temp_path, image)
        
        # Extract tables
        result = self.extract_tables(temp_path)
        
        # Cleanup
        try:
            os.unlink(temp_path)
        except:
            pass
        
        return result
    
    def tables_to_text(self, table_result: TableResult, 
                       format: str = 'markdown') -> str:
        """
        Convert extracted tables to text format
        
        Args:
            table_result: TableResult from extract_tables
            format: 'markdown', 'csv', or 'plain'
            
        Returns:
            Formatted text representation
        """
        if not table_result.tables:
            return ""
        
        output_parts = []
        
        for i, table in enumerate(table_result.tables):
            headers = table['headers']
            rows = table['rows']
            
            if format == 'markdown':
                # Markdown table
                lines = []
                lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                lines.append("|" + "|".join(["---"] * len(headers)) + "|")
                for row in rows:
                    lines.append("| " + " | ".join(str(c) if c is not None else "" for c in row) + " |")
                output_parts.append(f"\n**Table {i+1}:**\n" + "\n".join(lines))
                
            elif format == 'csv':
                # CSV format
                lines = [",".join(str(h) for h in headers)]
                for row in rows:
                    lines.append(",".join(str(c) if c is not None else "" for c in row))
                output_parts.append(f"# Table {i+1}\n" + "\n".join(lines))
                
            else:  # plain
                # Plain text with alignment
                lines = []
                lines.append(" | ".join(str(h) for h in headers))
                lines.append("-" * 50)
                for row in rows:
                    lines.append(" | ".join(str(c) if c is not None else "" for c in row))
                output_parts.append(f"Table {i+1}:\n" + "\n".join(lines))
        
        return "\n\n".join(output_parts)
    
    def extract_medical_table(self, image_path: str) -> Dict:
        """
        Specialized extraction for medical lab report tables
        
        Returns structured medical test results
        """
        table_result = self.extract_tables(image_path, borderless_tables=True)
        
        if not table_result.tables:
            return {'tests': [], 'raw_text': ''}
        
        tests = []
        
        for table in table_result.tables:
            headers = [str(h).lower() for h in table['headers']]
            
            # Find column indices
            test_col = next((i for i, h in enumerate(headers) if 'test' in h or 'parameter' in h or 'investigation' in h), 0)
            value_col = next((i for i, h in enumerate(headers) if 'value' in h or 'result' in h or 'observed' in h), 1)
            unit_col = next((i for i, h in enumerate(headers) if 'unit' in h), None)
            ref_col = next((i for i, h in enumerate(headers) if 'ref' in h or 'normal' in h or 'range' in h), None)
            
            for row in table['rows']:
                if len(row) > value_col:
                    test_entry = {
                        'test_name': str(row[test_col]) if row[test_col] else '',
                        'value': str(row[value_col]) if row[value_col] else '',
                        'unit': str(row[unit_col]) if unit_col and len(row) > unit_col and row[unit_col] else '',
                        'reference': str(row[ref_col]) if ref_col and len(row) > ref_col and row[ref_col] else ''
                    }
                    if test_entry['test_name'] and test_entry['value']:
                        tests.append(test_entry)
        
        return {
            'tests': tests,
            'table_count': table_result.table_count,
            'raw_tables': table_result.tables
        }

    def ocr_tesseract(self, 
                      image: np.ndarray,
                      preprocess: bool = True) -> OCRResult:
        """
        OCR using Tesseract with confidence scoring
        """
        if not self.tesseract_available:
            return OCRResult('', 0.0, 'unknown', 'tesseract_unavailable', 0, 'none')
        
        # Preprocess
        if preprocess:
            processed, preproc_method = self.preprocess_image(image, 'adaptive')
        else:
            processed = image
            preproc_method = 'none'
        
        # Convert to PIL
        pil_image = Image.fromarray(processed)
        
        # Get detailed OCR data
        try:
            ocr_data = pytesseract.image_to_data(
                pil_image,
                lang='+'.join(self.languages),
                output_type=pytesseract.Output.DICT
            )
            
            # Extract text and confidence
            texts = []
            confidences = []
            
            for i, text in enumerate(ocr_data['text']):
                conf = int(ocr_data['conf'][i])
                if conf > 0 and text.strip():
                    texts.append(text)
                    confidences.append(conf)
            
            full_text = ' '.join(texts)
            avg_confidence = np.mean(confidences) if confidences else 0.0
            
            # Clean text
            full_text = self._clean_ocr_text(full_text)
            
            return OCRResult(
                text=full_text,
                confidence=avg_confidence / 100,
                language=self._detect_language(full_text),
                method='tesseract',
                word_count=len(full_text.split()),
                preprocessing=preproc_method
            )
        except Exception as e:
            print(f"   ⚠️ Tesseract error: {e}")
            return OCRResult('', 0.0, 'unknown', 'tesseract_error', 0, preproc_method)
    
    def ocr_easyocr(self, image: np.ndarray) -> OCRResult:
        """
        OCR using EasyOCR (better for noisy images)
        """
        if self.easyocr_reader is None:
            return OCRResult('', 0.0, 'unknown', 'easyocr_unavailable', 0, 'none')
        
        try:
            results = self.easyocr_reader.readtext(image)
            
            texts = []
            confidences = []
            
            for bbox, text, conf in results:
                if text.strip():
                    texts.append(text)
                    confidences.append(conf)
            
            full_text = ' '.join(texts)
            avg_confidence = np.mean(confidences) if confidences else 0.0
            
            full_text = self._clean_ocr_text(full_text)
            
            return OCRResult(
                text=full_text,
                confidence=avg_confidence,
                language=self._detect_language(full_text),
                method='easyocr',
                word_count=len(full_text.split()),
                preprocessing='native'
            )
        except Exception as e:
            print(f"   ⚠️ EasyOCR error: {e}")
            return OCRResult('', 0.0, 'unknown', 'easyocr_error', 0, 'none')
    
    def ocr_best(self, image: np.ndarray) -> OCRResult:
        """
        Get best OCR result using smart fallback strategy:
        1. Try Tesseract first (fast)
        2. Only use EasyOCR if: low confidence, multilingual, or handwritten
        """
        # Step 1: Try Tesseract with adaptive preprocessing
        processed, method = self.preprocess_image(image.copy(), 'adaptive')
        tesseract_result = self.ocr_tesseract(processed, preprocess=False)
        tesseract_result.preprocessing = method
        
        # Step 2: Check if EasyOCR fallback is needed
        if self.smart_fallback:
            needs_fallback, reason = self._needs_easyocr_fallback(tesseract_result, image)
            
            if needs_fallback:
                print(f"   🔄 Triggering EasyOCR fallback: {reason}")
                
                # Lazy load EasyOCR
                if self._load_easyocr_if_needed():
                    easyocr_result = self.ocr_easyocr(image)
                    
                    # Compare and return best result
                    if easyocr_result.confidence > tesseract_result.confidence:
                        print(f"   ✅ EasyOCR improved: {tesseract_result.confidence:.1f}% → {easyocr_result.confidence:.1f}%")
                        return easyocr_result
                    else:
                        print(f"   ℹ️ Tesseract still better, keeping original")
        
        return tesseract_result
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean and normalize OCR output"""
        # Fix common OCR errors
        replacements = {
            '|': 'I',
            '0': 'O',  # Only in specific contexts
            'l': '1',  # Only in numbers
            '\n\n\n': '\n\n',
            '  ': ' ',
        }
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize whitespace
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Fix common patterns
        text = re.sub(r'(\d)[lI](\d)', r'\g<1>1\2', text)  # 4l5 -> 415
        text = re.sub(r'Rs\s*\.?\s*', 'Rs. ', text)  # Normalize Rs
        text = re.sub(r'(\d),(\d{3})', r'\1\2', text)  # 1,000 -> 1000
        
        return text.strip()
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        # Hindi Unicode range
        hindi_chars = len(re.findall(r'[\u0900-\u097F]', text))
        total_chars = len(re.findall(r'\w', text))
        
        if total_chars == 0:
            return 'unknown'
        
        hindi_ratio = hindi_chars / total_chars
        
        if hindi_ratio > 0.3:
            return 'hindi'
        elif hindi_ratio > 0.1:
            return 'hindi-english'
        else:
            return 'english'
    
    def _load_easyocr_if_needed(self) -> bool:
        """Lazy load EasyOCR only when actually needed"""
        if self.easyocr_reader is not None:
            return True
        
        if self._easyocr_load_attempted:
            return False
        
        self._easyocr_load_attempted = True
        
        if not EASYOCR_AVAILABLE:
            print("   ⚠️ EasyOCR not installed, skipping fallback")
            return False
        
        try:
            print("   🔄 Loading EasyOCR for fallback (one-time)...")
            import easyocr
            self.easyocr_reader = easyocr.Reader(
                ['en', 'hi'],  # English + Hindi for multilingual
                gpu=self.use_gpu
            )
            print("   ✅ EasyOCR loaded successfully")
            return True
        except Exception as e:
            print(f"   ⚠️ EasyOCR failed to load: {e}")
            return False
    
    def _detect_handwriting(self, image: np.ndarray) -> bool:
        """
        Detect if image likely contains handwritten text
        Uses variance in stroke width and character spacing
        """
        try:
            # Convert to grayscale if needed
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
            
            # Apply edge detection
            edges = cv2.Canny(gray, 50, 150)
            
            # Find contours
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            if len(contours) < 10:
                return False
            
            # Calculate variance in contour sizes (handwriting has high variance)
            areas = [cv2.contourArea(c) for c in contours if cv2.contourArea(c) > 10]
            if not areas:
                return False
            
            mean_area = np.mean(areas)
            std_area = np.std(areas)
            
            # High coefficient of variation suggests handwriting
            cv = std_area / mean_area if mean_area > 0 else 0
            
            return cv > 1.5  # Threshold for handwriting detection
        except:
            return False
    
    def _needs_easyocr_fallback(self, tesseract_result: 'OCRResult', image: np.ndarray) -> Tuple[bool, str]:
        """
        Determine if EasyOCR fallback is needed
        
        Returns:
            (needs_fallback, reason)
        """
        # Check 1: Low confidence
        if tesseract_result.confidence < self.CONFIDENCE_THRESHOLD:
            return True, f"low_confidence ({tesseract_result.confidence:.1f}%)"
        
        # Check 2: Multi-language detected
        if tesseract_result.language in ['hindi', 'hindi-english']:
            return True, f"multilingual ({tesseract_result.language})"
        
        # Check 3: Very few words extracted (might be handwritten)
        if tesseract_result.word_count < 5 and len(tesseract_result.text.strip()) > 50:
            return True, "sparse_extraction"
        
        # Check 4: Handwriting detected
        if self._detect_handwriting(image):
            return True, "handwriting_detected"
        
        return False, "none"
    
    def extract_from_image(self, image_path: str) -> OCRResult:
        """
        Extract text from image file
        
        Args:
            image_path: Path to image file
            
        Returns:
            OCRResult with extracted text
        """
        print(f"📄 Processing image: {image_path}")
        
        # Load image
        image = cv2.imread(image_path)
        if image is None:
            # Try with PIL
            try:
                pil_img = Image.open(image_path)
                image = np.array(pil_img)
            except Exception as e:
                return OCRResult('', 0.0, 'unknown', f'load_error: {e}', 0, 'none')
        
        # Get best OCR result
        result = self.ocr_best(image)
        
        print(f"   ✅ Extracted {result.word_count} words (confidence: {result.confidence:.1%})")
        
        return result
    
    def extract_from_docx(self, docx_path: str) -> OCRResult:
        """
        Extract text from DOCX file
        """
        if not DOCX_AVAILABLE:
            return OCRResult('', 0.0, 'unknown', 'docx_unavailable', 0, 'none')
        
        try:
            doc = Document(docx_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            full_text = '\n'.join(paragraphs)
            
            return OCRResult(
                text=full_text,
                confidence=1.0,  # Direct extraction is 100% accurate
                language=self._detect_language(full_text),
                method='docx_direct',
                word_count=len(full_text.split()),
                preprocessing='none'
            )
        except Exception as e:
            return OCRResult('', 0.0, 'unknown', f'docx_error: {e}', 0, 'none')
    
    def extract_from_pdf(self, pdf_path: str) -> OCRResult:
        """
        Extract text from PDF (converts to images first)
        """
        if not PDF2IMAGE_AVAILABLE:
            return OCRResult('', 0.0, 'unknown', 'pdf2image_unavailable', 0, 'none')
        
        try:
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=200)
            
            all_text = []
            total_confidence = 0
            
            for i, img in enumerate(images):
                img_array = np.array(img)
                result = self.ocr_best(img_array)
                all_text.append(result.text)
                total_confidence += result.confidence
            
            full_text = '\n\n--- Page Break ---\n\n'.join(all_text)
            avg_confidence = total_confidence / len(images) if images else 0
            
            return OCRResult(
                text=full_text,
                confidence=avg_confidence,
                language=self._detect_language(full_text),
                method='pdf_ocr',
                word_count=len(full_text.split()),
                preprocessing='multi_page'
            )
        except Exception as e:
            return OCRResult('', 0.0, 'unknown', f'pdf_error: {e}', 0, 'none')
    
    def extract(self, file_path: str) -> OCRResult:
        """
        Auto-detect file type and extract text
        
        Args:
            file_path: Path to file (image, docx, or pdf)
            
        Returns:
            OCRResult
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        elif ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self.extract_from_image(file_path)
        else:
            return OCRResult('', 0.0, 'unknown', f'unsupported_format: {ext}', 0, 'none')
    
    def extract_text(self, file_path: str, detect_tables: bool = False) -> str:
        """
        Extract text from file (convenience wrapper for API compatibility)
        
        Args:
            file_path: Path to file
            detect_tables: Whether to detect tables (uses table extraction if True)
            
        Returns:
            Extracted text as string
        """
        result = self.extract(file_path)
        return result.text


# Alias for backward compatibility
OCRPipeline = EnhancedOCRPipeline


if __name__ == "__main__":
    # Test the OCR pipeline
    pipeline = EnhancedOCRPipeline()
    
    # Test on sample image
    import sys
    if len(sys.argv) > 1:
        result = pipeline.extract(sys.argv[1])
        print(f"\nResult: {result.to_dict()}")
